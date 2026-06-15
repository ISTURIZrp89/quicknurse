import os
import tempfile
from app.services.text_extractor import extract_text


def test_extract_markdown():
    content = "# RCP Avanzado\n\nLa reanimación cardiopulmonar avanzada incluye:\n- Vía aérea\n- Desfibrilación\n- Fármacos"
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", delete=False, encoding="utf-8") as f:
        f.write(content)
        path = f.name
    try:
        result = extract_text(path)
        assert "reanimación cardiopulmonar" in result
        assert "Desfibrilación" in result
    finally:
        os.unlink(path)


def test_extract_txt():
    content = "Signos vitales normales: FC 60-100, FR 12-20, PA 120/80"
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8") as f:
        f.write(content)
        path = f.name
    try:
        result = extract_text(path)
        assert result == content
    finally:
        os.unlink(path)


def test_extract_pdf():
    try:
        import fitz
    except ImportError:
        pytest.skip("PyMuPDF not installed")
    import pytest

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Protocolo de sepsis: administrar antibióticos dentro de la primera hora.")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        path = f.name
    doc.save(path)
    doc.close()

    try:
        result = extract_text(path)
        assert "sepsis" in result.lower()
        assert "antibióticos" in result.lower()
    finally:
        os.unlink(path)


def test_extract_docx():
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")
    import pytest

    doc = Document()
    doc.add_paragraph("Escala de Glasgow:")
    doc.add_paragraph("Ocular: 1-4")
    doc.add_paragraph("Verbal: 1-5")
    doc.add_paragraph("Motor: 1-6")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    doc.save(path)

    try:
        result = extract_text(path)
        assert "Glasgow" in result
        assert "Ocular" in result
        assert "Verbal" in result
    finally:
        os.unlink(path)


def test_extract_unknown_extension():
    with tempfile.NamedTemporaryFile(suffix=".xyz", mode="w", delete=False) as f:
        f.write("test content")
        path = f.name
    try:
        result = extract_text(path)
        assert result == ""
    finally:
        os.unlink(path)
